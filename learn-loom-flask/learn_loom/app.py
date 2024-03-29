#imports

from io import BytesIO 
#importing BytesIO class from IO module

from bs4 import BeautifulSoup # Importing BeautifulSoup class from bs4 module
from docx import Document
from flask import Flask, request, redirect, flash, jsonify, url_for, send_file, make_response
from flask_cors import CORS
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
from werkzeug.utils import secure_filename
from modules import llm_service, video_indexer_service
from db import video_repository
from util import constant, util

import os
import uuid

#Creating the flask application
app = Flask(_name_)
CORS(app)

# This is the path to the directory to save the uploaded files.
UPLOAD_FOLDER = constant.BASE_DIR + '\\uploaded-files'
THUMBNAIL_FOLDER = UPLOAD_FOLDER + '\\thumb'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['THUMBNAIL_FOLDER'] = THUMBNAIL_FOLDER

ACCOUNT_ID = '37df8c3e-bd8c-476e-82fb-86fd52f1f0a0'
SUBSCRIPTION_KEY = 'a9a8ff314bff420faa39c3237ffc5be9'

#Gets infromation about the videos in database
@app.route('/api/video', methods=['GET'])
def get_all_videos():
    conn = video_repository.create_connection()
    with conn:
        list_of_videos = video_repository.read_all_videos(conn)
        dict_of_videos = [vid.to_dict() for vid in list_of_videos]
        print('list_of_videos', dict_of_videos)
        return jsonify(dict_of_videos)

#Gets information about single video by its uuid
@app.route('/api/video/<video_uuid>', methods=['GET'])
def get_video_by_uuid(video_uuid):
    conn = video_repository.create_connection()
    with conn:
        video = video_repository.read_video_by_video_uuid(conn, video_uuid)
        video_indexer_id = video.video_indexer_id
        dict_of_videos = video.to_dict()
        print('dict_of_videos', dict_of_videos)
        return jsonify(dict_of_videos)

#Generates PDF dcument
@app.route('/api/video/<video_uuid>/pdf', methods=['GET'])
def get_video_content_as_pdf_by_uuid(video_uuid):
    type = request.args.get('type')

    conn = video_repository.create_connection()
    with conn:
        video = video_repository.read_video_by_video_uuid(conn, video_uuid)
        video_indexer_id = video.video_indexer_id
        dict_of_videos = video.to_dict()
        print('dict_of_videos', dict_of_videos)

        content = ''
        name = ''

        if type == 'questions':
            content = video.question
            name = 'Questions'
        elif type == 'flashCards':
            content = video.flash_card
            name = 'Flash Cards'
        else:
            content = video.note
            name = 'Notes'

        ##pdf = pdfkit.from_string(content, False)
        ##pdf = HTML(string=content).write_pdf()

        html_content = """
            <html>
            <head>
                <title>PDF Example</title>
                <style>
                    body {
                        font-family: Arial, sans-serif;
                        color: #333;
                    }
                    h1 {
                        color: #333;
                        font-size: 24px;
                        margin-bottom: 10px;
                    }
                    p {
                        color: #666;
                        font-size: 16px;
                        margin-bottom: 10px;
                    }
                </style>
            </head>
            <body>
                <h1>Hello, PDF!</h1>
                <pre><p>This is a sample PDF generated from HTML string.</p></pre>
            </body>
            </html>
            """

        response_buffer = BytesIO()
        pdf = SimpleDocTemplate(response_buffer, pagesize=letter)

        styles = getSampleStyleSheet()
        style_normal = styles["Normal"]
        style_heading = styles["Heading1"]

        soup = BeautifulSoup(html_content, 'html.parser')
        pdf_content = []
        for tag in soup.find_all():
            if tag.name == 'h1':
                pdf_content.append(Paragraph(name, style_heading))
            elif tag.name == 'p':
                pdf_content.append(Paragraph(content, style_normal))

        pdf.build(pdf_content)

        response_buffer.seek(0)
        response = make_response(response_buffer.getvalue())
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = 'inline; filename=output.pdf'

        return response

#Generates word document
@app.route('/api/video/<video_uuid>/word-doc', methods=['GET'])
def get_video_content_as_word_by_uuid(video_uuid):
    type = request.args.get('type')

    conn = video_repository.create_connection()
    with conn:
        video = video_repository.read_video_by_video_uuid(conn, video_uuid)
        video_indexer_id = video.video_indexer_id
        dict_of_videos = video.to_dict()
        print('dict_of_videos', dict_of_videos)

        content = ''
        name = ''

        if type == 'questions':
            content = video.question
            name = 'Questions'
        elif type == 'flashCards':
            content = video.flash_card
            name = 'Flash Cards'
        else:
            content = video.note
            name = 'Notes'

        doc = Document()

        # Add content to the document
        doc.add_heading(name, level=1)
        doc.add_paragraph(content)

        # Save the document to a file
        doc.save('sample.docx')

        # Send the document as a response
        return send_file('sample.docx', as_attachment=True)


#Gets a thumbnail for the video to display
@app.route('/api/video/<video_uuid>/thumb', methods=['GET'])
def get_video_thumbnail(video_uuid):
    return send_file(os.path.join(app.config["THUMBNAIL_FOLDER"], f"{video_uuid}.jpg"), mimetype='image/jpeg')


@app.route('/api/video-upload', methods=['POST'])
def upload_file():
    print('invoking video-upload')
    # Check if the post request has the file part
    if 'file' not in request.files:
        flash('No file part')
        print('No file part')
        return redirect(request.url)
    file = request.files['file']

    if file.filename == '':
        flash('No selected file')
        print('No selected file')
        return redirect(request.url)

    if file and util.allowed_file(file.filename):
        original_video_name = file.filename
        video_uuid = str(uuid.uuid4())
        video_ext = os.path.splitext(file.filename)[1]
        video_name_with_ext = video_uuid + video_ext
        file_path = UPLOAD_FOLDER + '/' + video_name_with_ext
        if os.path.exists(file_path):
            os.remove(file_path)

        upload_video = secure_filename(video_name_with_ext)
        upload_video_path = os.path.join(app.config['UPLOAD_FOLDER'], upload_video)
        file.save(upload_video_path)
        thumbnail_path = os.path.join(app.config['THUMBNAIL_FOLDER'], f"{video_uuid}.jpg")
        util.generate_thumbnail(upload_video_path, thumbnail_path)

        conn = video_repository.create_connection()
        with conn:
            video_id = video_repository.create_base_video(conn, video_uuid, original_video_name, video_ext)
            response = jsonify(
                {'video_id': video_id, 'video_uuid': video_uuid, 'original_video_name': original_video_name,
                 'video_ext': video_ext})
            return response


@app.route('/api/process-video', methods=['POST'])
def video_process():
    print('process-video')

    video = request.get_json().get('video_uuid')
    video_uuid = video['video_uuid']
    video_ext = video['video_ext']
    print('video_uuid', video_uuid)

    video_path = UPLOAD_FOLDER + '/' + video_uuid + video_ext

    # Get access token
    access_token = video_indexer_service.get_access_token(ACCOUNT_ID, SUBSCRIPTION_KEY)

    # Upload video and get video ID
    video_indexer_id = video_indexer_service.upload_video(access_token, ACCOUNT_ID, video_path, video_uuid)['id']
    conn = video_repository.create_connection()
    with conn:
        video_id = video_repository.update_video_indexer_id(conn, video_indexer_id, video_uuid)
        print('done - [/api/process-video]')
        return jsonify({'video_uuid': video_id, 'video_indexer_id': video_indexer_id})


@app.route('/api/process', methods=['POST'])
def process():
    print('process-text')

    video = request.get_json().get('video_uuid')
    video_uuid = video['video_uuid']

    conn = video_repository.create_connection()
    with conn:
        video = video_repository.read_video_by_video_uuid(conn, video_uuid)
        video_indexer_id = video.video_indexer_id

        # Get access token
        access_token = video_indexer_service.get_access_token(ACCOUNT_ID, SUBSCRIPTION_KEY)

        VideoIndex = video_indexer_service.get_video_index(access_token, ACCOUNT_ID, video_indexer_id)
        transcript = VideoIndex.transcript
        print('transcript:', transcript)

        inputs = [
            'Create a note with topics from the embedded text by only including the relvant and important educational content of around 500 words.',
            'Create 10 questions from the text that are educationally relevant, also include ideal answers.',
            'Create 10 flashcards from the most educationally relevant content given.']

        note = llm_service.invoke_llm(inputs[0] + transcript)
        questions = llm_service.invoke_llm(inputs[1] + transcript)
        flash_cards = llm_service.invoke_llm(inputs[2] + transcript)

        conn = video_repository.create_connection()
        with conn:
            video_id = video_repository.update_video_content(conn, note, questions, flash_cards, video_uuid)
            return jsonify({'state': VideoIndex.state, 'video_id': video_id, 'note': note, 'questions': questions,
                            'flash_cards': flash_cards})
